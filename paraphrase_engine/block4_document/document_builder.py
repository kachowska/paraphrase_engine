"""
Block 4: Document Builder
Automatically replaces original fragments with paraphrased ones in .docx documents
"""

import os
import logging
from typing import List, Tuple, Optional
from pathlib import Path
import shutil
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_UNDERLINE

from ..block5_logging.logger import SystemLogger

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Handles document manipulation and fragment replacement"""
    
    def __init__(self):
        self.system_logger = SystemLogger()
    
    async def replace_fragments(
        self,
        source_file_path: str,
        output_file_path: str,
        original_fragments: List[str],
        paraphrased_fragments: List[str],
        progress_callback=None
    ) -> bool:
        """
        Replace fragments in document
        CRITICAL: Replacements are made in reverse order to preserve text indexing
        
        Args:
            source_file_path: Path to source .docx file
            output_file_path: Path for output .docx file
            original_fragments: List of original text fragments
            paraphrased_fragments: List of paraphrased text fragments
            
        Returns:
            bool: True if successful, False otherwise
        """
        
        if len(original_fragments) != len(paraphrased_fragments):
            logger.error("Mismatch between original and paraphrased fragments count")
            return False
        
        try:
            # Create a copy of the source file
            shutil.copy2(source_file_path, output_file_path)
            logger.info(f"Created document copy: {output_file_path}")
            
            # Open the document
            doc = Document(output_file_path)
            
            # Process replacements in REVERSE order (critical requirement)
            replacements_made = 0
            skipped_fragments = []
            total_fragments = len(original_fragments)
            
            for i in range(len(original_fragments) - 1, -1, -1):
                original = original_fragments[i]
                paraphrased = paraphrased_fragments[i]
                fragment_number = i + 1
                
                logger.info(f"Processing fragment {fragment_number}/{total_fragments} (reverse order)")
                
                # Send progress update every 20 fragments or at milestones
                if progress_callback and (fragment_number % 20 == 0 or fragment_number in [total_fragments, total_fragments // 2, total_fragments // 4]):
                    progress_percent = int(((total_fragments - i) / total_fragments) * 100)
                    try:
                        await progress_callback(
                            f"📝 *Замена в документе*\n\n"
                            f"🔍 Обработано: {total_fragments - i}/{total_fragments} фрагментов\n"
                            f"📈 {progress_percent}%\n"
                            f"✅ Найдено и заменено: {replacements_made}"
                        )
                    except Exception as e:
                        logger.warning(f"Error sending progress update during document replacement: {e}")
                
                # Try to replace the fragment
                replaced = await self._replace_fragment_in_document(
                    doc,
                    original,
                    paraphrased,
                    fragment_index=i
                )
                
                if replaced:
                    replacements_made += 1
                    await self.system_logger.log_fragment_replaced(
                        fragment_index=i,
                        original_length=len(original),
                        paraphrased_length=len(paraphrased)
                    )
                else:
                    skipped_fragments.append(i)
                    await self.system_logger.log_fragment_not_found(
                        fragment_index=i,
                        fragment_text=original[:50] + "..." if len(original) > 50 else original
                    )
            
            # Save the modified document
            doc.save(output_file_path)
            
            logger.info(
                f"Document processing complete. "
                f"Replacements: {replacements_made}/{len(original_fragments)}, "
                f"Skipped: {len(skipped_fragments)}"
            )
            
            # Log completion
            await self.system_logger.log_document_processed(
                source_path=source_file_path,
                output_path=output_file_path,
                total_fragments=len(original_fragments),
                replaced_fragments=replacements_made,
                skipped_fragments=len(skipped_fragments)
            )
            
            return True
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            await self.system_logger.log_error(
                chat_id=0,
                operation="document_builder",
                error_message=str(e)
            )
            return False
    
    async def _replace_fragment_in_document(
        self,
        doc: DocumentType,
        original_text: str,
        replacement_text: str,
        fragment_index: int
    ) -> bool:
        """
        Replace a single fragment in the document
        
        Returns:
            bool: True if replacement was made, False if fragment not found
        """
        
        # Normalize texts for comparison (soft normalization)
        original_normalized = self._normalize_text(original_text)
        
        # Track if replacement was made
        replacement_made = False
        
        # Search through all paragraphs (including multi-paragraph search)
        paragraphs_list = list(doc.paragraphs)
        
        for i, paragraph in enumerate(paragraphs_list):
            paragraph_text = paragraph.text
            paragraph_normalized = self._normalize_text(paragraph_text)
            
            # Check if the fragment exists in this paragraph (exact match preferred)
            # First try exact match
            if original_text in paragraph_text:
                success = self._replace_in_paragraph_with_formatting(
                    paragraph,
                    original_text,
                    replacement_text
                )
                
                if success:
                    replacement_made = True
                    logger.info(f"Replaced fragment {fragment_index} in paragraph (exact match)")
                    break
            
            # If exact match failed, try normalized match in single paragraph
            elif original_normalized in paragraph_normalized:
                # Try to find the actual text in the paragraph (accounting for whitespace differences)
                actual_text = self._find_actual_text_in_paragraph(paragraph, original_normalized)
                
                if actual_text:
                    success = self._replace_in_paragraph_with_formatting(
                        paragraph,
                        actual_text,
                        replacement_text
                    )
                    
                    if success:
                        replacement_made = True
                        logger.info(f"Replaced fragment {fragment_index} in paragraph (normalized match)")
                        break
            
            # If still not found, try searching across multiple paragraphs (for fragments split across lines)
            if not replacement_made and i < len(paragraphs_list) - 1:
                # Combine current paragraph with next paragraphs (up to 5 paragraphs for long fragments)
                max_combine = min(5, len(paragraphs_list) - i)
                combined_text = paragraph_text
                combined_normalized = paragraph_normalized
                combined_paragraphs = [paragraph]
                
                for j in range(1, max_combine):
                    next_para = paragraphs_list[i + j]
                    combined_text += " " + next_para.text
                    combined_normalized += " " + self._normalize_text(next_para.text)
                    combined_paragraphs.append(next_para)
                
                # Check normalized match in combined text
                if original_normalized in combined_normalized:
                    # Try to find the actual text across multiple paragraphs
                    actual_text = self._find_actual_text_across_paragraphs(combined_paragraphs, original_normalized)
                    
                    if actual_text:
                        # Replace in the first paragraph that contains the start of the fragment
                        success = self._replace_in_paragraph_with_formatting(
                            paragraph,
                            actual_text,
                            replacement_text
                        )
                        
                        if success:
                            replacement_made = True
                            logger.info(f"Replaced fragment {fragment_index} in paragraph (multi-paragraph normalized match, {len(combined_paragraphs)} paragraphs)")
                            break
        
        # Also check in tables
        if not replacement_made:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text
                            paragraph_normalized = self._normalize_text(paragraph_text)
                            
                            # Try exact match first
                            if original_text in paragraph_text:
                                success = self._replace_in_paragraph_with_formatting(
                                    paragraph,
                                    original_text,
                                    replacement_text
                                )
                                
                                if success:
                                    replacement_made = True
                                    logger.info(f"Replaced fragment {fragment_index} in table (exact match)")
                                    break
                            
                            # Try normalized match
                            elif original_normalized in paragraph_normalized:
                                actual_text = self._find_actual_text_in_paragraph(paragraph, original_normalized)
                                
                                if actual_text:
                                    success = self._replace_in_paragraph_with_formatting(
                                        paragraph,
                                        actual_text,
                                        replacement_text
                                    )
                                    
                                    if success:
                                        replacement_made = True
                                        logger.info(f"Replaced fragment {fragment_index} in table (normalized match)")
                                        break
                        
                        if replacement_made:
                            break
                    
                    if replacement_made:
                        break
                
                if replacement_made:
                    break
        
        # If still not found, try keyword-based search (last resort)
        if not replacement_made:
            # Extract meaningful words from the fragment (remove common words, numbers, short words)
            import re
            words = re.findall(r'\b\w{4,}\b', original_normalized)  # Words with 4+ characters
            if len(words) >= 3:  # Need at least 3 meaningful words
                # Search for paragraphs containing at least 2 of these words
                for paragraph in doc.paragraphs:
                    para_normalized = self._normalize_text(paragraph.text)
                    matching_words = sum(1 for word in words if word.lower() in para_normalized.lower())
                    
                    if matching_words >= 2:  # At least 2 words match
                        # Try to find a substring that contains these words
                        # Use the longest matching substring
                        best_match = self._find_best_keyword_match(paragraph, words, original_normalized)
                        
                        if best_match:
                            success = self._replace_in_paragraph_with_formatting(
                                paragraph,
                                best_match,
                                replacement_text
                            )
                            
                            if success:
                                replacement_made = True
                                logger.info(f"Replaced fragment {fragment_index} in paragraph (keyword-based match, {matching_words}/{len(words)} words)")
                                break
                
                # If still not found, try searching across multiple paragraphs with keywords
                if not replacement_made and len(words) >= 5:
                    for i in range(len(paragraphs_list) - 1):
                        # Combine up to 3 paragraphs
                        combined_para_text = paragraphs_list[i].text
                        combined_para_normalized = self._normalize_text(combined_para_text)
                        combined_paras = [paragraphs_list[i]]
                        
                        for j in range(1, min(3, len(paragraphs_list) - i)):
                            next_para = paragraphs_list[i + j]
                            combined_para_text += " " + next_para.text
                            combined_para_normalized += " " + self._normalize_text(next_para.text)
                            combined_paras.append(next_para)
                        
                        matching_words = sum(1 for word in words if word.lower() in combined_para_normalized.lower())
                        if matching_words >= 3:  # At least 3 words match in combined text
                            # Try to find text across paragraphs
                            best_match = self._find_actual_text_across_paragraphs(combined_paras, original_normalized)
                            
                            if best_match:
                                success = self._replace_in_paragraph_with_formatting(
                                    paragraphs_list[i],
                                    best_match,
                                    replacement_text
                                )
                                
                                if success:
                                    replacement_made = True
                                    logger.info(f"Replaced fragment {fragment_index} across {len(combined_paras)} paragraphs (multi-paragraph keyword match, {matching_words}/{len(words)} words)")
                                    break
                        
                        if replacement_made:
                            break
        
        if not replacement_made:
            # Enhanced logging to help debug why fragment wasn't found
            logger.warning(
                f"Fragment {fragment_index} not found in document. "
                f"Original: {original_text[:100]}..."
            )
            logger.debug(
                f"Fragment {fragment_index} normalized: {original_normalized[:150]}..."
            )
            # Log first few paragraphs for debugging
            if len(doc.paragraphs) > 0:
                for para_idx, para in enumerate(doc.paragraphs[:3]):
                    para_text = para.text[:200]
                    para_normalized = self._normalize_text(para.text)[:200]
                    logger.debug(
                        f"Paragraph {para_idx} sample: {para_text}... "
                        f"(normalized: {para_normalized}...)"
                    )
                    # Check if any words from fragment are in this paragraph
                    fragment_words = set(original_normalized.lower().split()[:5])  # First 5 words
                    para_words = set(para_normalized.lower().split())
                    common_words = fragment_words.intersection(para_words)
                    if common_words:
                        logger.debug(
                            f"Paragraph {para_idx} shares {len(common_words)} words with fragment: {list(common_words)[:3]}"
                        )
        
        return replacement_made
    
    def _find_actual_text_across_paragraphs(
        self,
        paragraphs: List[Paragraph],
        normalized_search: str
    ) -> Optional[str]:
        """
        Find actual text across multiple paragraphs that matches normalized search string
        """
        # Combine all paragraph texts
        combined_text = " ".join(para.text for para in paragraphs)
        combined_normalized = self._normalize_text(combined_text)
        
        # Check if normalized search is in combined text
        if normalized_search not in combined_normalized:
            return None
        
        # Find position in normalized text
        start_pos = combined_normalized.find(normalized_search)
        if start_pos == -1:
            return None
        
        # Map normalized position back to actual text
        # Count normalized characters to find approximate position
        norm_count = 0
        actual_start = 0
        actual_end = len(combined_text)
        
        for i, char in enumerate(combined_text):
            char_norm = self._normalize_text(char)
            if char_norm:
                if norm_count == start_pos:
                    actual_start = i
                if norm_count == start_pos + len(normalized_search):
                    actual_end = i
                    break
                norm_count += 1
        
        # Extract candidate text
        candidate = combined_text[actual_start:actual_end]
        
        # Verify it normalizes correctly
        if self._normalize_text(candidate) == normalized_search:
            return candidate
        
        # If exact match failed, try word-based matching
        search_words = normalized_search.split()
        if len(search_words) < 3:
            return None
        
        # Find first 3 words in combined text
        first_words = " ".join(search_words[:3])
        if first_words in combined_normalized:
            first_pos = combined_normalized.find(first_words)
            # Try to extract text starting from first words
            norm_count = 0
            actual_start = 0
            for i, char in enumerate(combined_text):
                char_norm = self._normalize_text(char)
                if char_norm:
                    if norm_count == first_pos:
                        actual_start = i
                        break
                    norm_count += 1
            
            # Extract a reasonable amount of text (at least the search length)
            min_length = len(normalized_search)
            actual_end = min(actual_start + min_length * 2, len(combined_text))
            candidate = combined_text[actual_start:actual_end]
            
            # Verify it contains the key words
            candidate_normalized = self._normalize_text(candidate)
            if all(word in candidate_normalized for word in search_words[:3]):
                return candidate
        
        return None
    
    def _replace_in_paragraph_with_formatting(
        self,
        paragraph: Paragraph,
        original_text: str,
        replacement_text: str
    ) -> bool:
        """
        Replace text in paragraph while preserving formatting
        
        This is a complex operation that tries to maintain the original formatting
        Handles cases where text spans multiple runs
        """
        
        try:
            # Get full paragraph text
            full_text = paragraph.text
            
            # Check if original text exists (try exact match first)
            start_index = full_text.find(original_text)
            
            if start_index == -1:
                # Try normalized match
                normalized_full = self._normalize_text(full_text)
                normalized_original = self._normalize_text(original_text)
                
                if normalized_original not in normalized_full:
                    return False
                
                # Find normalized positions
                norm_start = normalized_full.find(normalized_original)
                norm_end = norm_start + len(normalized_original)
                
                # Build mapping from normalized positions to actual positions
                # by simulating the normalization process
                norm_to_actual = {}  # Maps normalized index to actual index
                norm_pos = 0
                in_whitespace = False
                
                for actual_pos, char in enumerate(full_text):
                    if not char.isspace():
                        # Non-whitespace character: always contributes to normalized text
                        norm_to_actual[norm_pos] = actual_pos
                        norm_pos += 1
                        in_whitespace = False
                    else:
                        # Whitespace: only first whitespace in sequence contributes
                        if not in_whitespace:
                            norm_to_actual[norm_pos] = actual_pos
                            norm_pos += 1
                            in_whitespace = True
                        # Subsequent whitespace doesn't contribute to normalized text
                
                # Map normalized start and end to actual positions
                if norm_start in norm_to_actual:
                    start_index = norm_to_actual[norm_start]
                else:
                    return False
                
                # For end position, find the actual position corresponding to norm_end
                # If exact match not found, use the position after the last character
                # that contributes to normalized text up to norm_end
                if norm_end in norm_to_actual:
                    end_index = norm_to_actual[norm_end]
                else:
                    # Find the last actual position that maps to a normalized position < norm_end
                    # and set end_index to the position after it
                    max_actual = -1
                    for norm_idx, actual_idx in norm_to_actual.items():
                        if norm_idx < norm_end:
                            max_actual = max(max_actual, actual_idx)
                    if max_actual >= 0:
                        # Find the end of the whitespace sequence or next non-whitespace
                        end_index = max_actual + 1
                        # Skip any trailing whitespace after this position
                        while end_index < len(full_text) and full_text[end_index].isspace():
                            end_index += 1
                    else:
                        end_index = len(full_text)
            else:
                end_index = start_index + len(original_text)
            
            # Build new paragraph content
            new_runs = []
            current_pos = 0
            replacement_added = False
            
            for run in paragraph.runs:
                run_text = run.text
                run_start = current_pos
                run_end = current_pos + len(run_text)
                
                # Case 1: Run is completely before the replacement area
                if run_end <= start_index:
                    new_runs.append((run_text, run))
                
                # Case 2: Run is completely after the replacement area
                elif run_start >= end_index:
                    new_runs.append((run_text, run))
                
                # Case 3: Run contains the start of replacement area
                elif run_start < start_index and run_end > start_index:
                    # Part before replacement
                    before_text = run_text[:start_index - run_start]
                    if before_text:
                        new_runs.append((before_text, run))
                    
                    # Add replacement text with same formatting (only once)
                    if not replacement_added:
                        new_runs.append((replacement_text, run))
                        replacement_added = True
                    
                    # Check if replacement ends in this run
                    if run_end >= end_index:
                        # Part after replacement
                        after_text = run_text[end_index - run_start:]
                        if after_text:
                            new_runs.append((after_text, run))
                
                # Case 4: Run is completely within replacement area
                elif run_start >= start_index and run_end <= end_index:
                    # Skip this run (it's being replaced)
                    # But add replacement if we haven't yet
                    if not replacement_added:
                        new_runs.append((replacement_text, run))
                        replacement_added = True
                
                # Case 5: Run contains the end of replacement area
                elif run_start < end_index and run_end > end_index:
                    # Add replacement if we haven't yet
                    if not replacement_added:
                        new_runs.append((replacement_text, run))
                        replacement_added = True
                    
                    # Part after replacement
                    after_text = run_text[end_index - run_start:]
                    if after_text:
                        new_runs.append((after_text, run))
                
                current_pos = run_end
            
            # If replacement wasn't added (edge case), add it at the start position
            if not replacement_added:
                # Find the run that contains start_index
                current_pos = 0
                for i, run in enumerate(paragraph.runs):
                    run_start = current_pos
                    run_end = current_pos + len(run.text)
                    if run_start <= start_index < run_end:
                        # Insert replacement before this run
                        new_runs.insert(i, (replacement_text, run))
                        break
                    current_pos = run_end
            
            # Clear paragraph and rebuild with new runs
            paragraph.clear()
            
            for text, original_run in new_runs:
                if text:  # Only add non-empty runs
                    new_run = paragraph.add_run(text)
                    # Copy formatting from original run
                    self._copy_run_formatting(original_run, new_run)
            
            return True
            
        except Exception as e:
            logger.error(f"Error replacing text with formatting: {e}", exc_info=True)
            
            # Fallback: Simple replacement without formatting preservation
            try:
                full_text = paragraph.text
                paragraph.text = full_text.replace(original_text, replacement_text)
                logger.warning(f"Used fallback replacement method for paragraph")
                return True
            except Exception as fallback_error:
                logger.error(f"Fallback replacement also failed: {fallback_error}")
                return False
    
    def _copy_run_formatting(self, source_run: Run, target_run: Run):
        """Copy formatting from source run to target run"""
        try:
            if source_run.bold is not None:
                target_run.bold = source_run.bold
            if source_run.italic is not None:
                target_run.italic = source_run.italic
            if source_run.underline is not None and source_run.underline != WD_UNDERLINE.NONE:
                # underline может быть bool или WD_UNDERLINE enum
                target_run.underline = source_run.underline  # type: ignore
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            if source_run.font.highlight_color:
                target_run.font.highlight_color = source_run.font.highlight_color
        except Exception as e:
            logger.debug(f"Could not copy some formatting: {e}")
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize text for comparison
        Handles different whitespace, line breaks, etc.
        Uses soft normalization to preserve text structure
        Removes PDF artifacts like page numbers and citations
        """
        import re
        # Remove PDF artifacts: page numbers, citations like [39, c. 126], [14], etc.
        # Pattern: [number, c. number] or [number] or "с. number" or standalone numbers
        text = re.sub(r'\[\d+[,\s]*(?:[сc]\.\s*)?\d*\]', '', text)  # [39, c. 126] or [14]
        text = re.sub(r'[сc]\.\s*\d+', '', text)  # с. 51
        text = re.sub(r'^\d+\s+\d+\s*', '', text)  # "61 19 января" -> "19 января"
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)  # Standalone numbers on lines
        
        # Normalize line endings first
        text = text.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
        # Replace multiple spaces/tabs with single space (but preserve single spaces)
        text = re.sub(r'[ \t]+', ' ', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text
    
    def _find_actual_text_in_paragraph(self, paragraph: Paragraph, normalized_search: str) -> Optional[str]:
        """
        Find the actual text in paragraph that matches normalized search string
        Accounts for whitespace differences while preserving the actual text structure
        """
        # Get all text from runs
        full_text = paragraph.text
        normalized_full = self._normalize_text(full_text)
        
        # Find position in normalized text
        start_pos = normalized_full.find(normalized_search)
        if start_pos == -1:
            return None
        
        # Try to find corresponding position in actual text
        # This is approximate - we'll use a sliding window approach
        search_length = len(normalized_search)
        
        # Build normalized text with position mapping
        normalized_chars = []
        actual_chars = []
        pos = 0
        
        for run in paragraph.runs:
            for char in run.text:
                actual_chars.append((char, pos))
                normalized_char = self._normalize_text(char)
                if normalized_char:
                    normalized_chars.append((normalized_char, pos))
                pos += 1
        
        # Find start and end positions
        normalized_text = ''.join([c[0] for c in normalized_chars])
        search_start = normalized_text.find(normalized_search)
        
        if search_start == -1:
            return None
        
        # Map back to actual text positions
        # This is complex, so we'll use a simpler approach:
        # Extract a window from the actual text that should contain our fragment
        actual_start = 0
        actual_end = len(full_text)
        
        # Count normalized characters to find approximate position
        norm_count = 0
        for i, char in enumerate(full_text):
            norm_char = self._normalize_text(char)
            if norm_char:
                if norm_count == search_start:
                    actual_start = i
                if norm_count == search_start + search_length:
                    actual_end = i
                    break
                norm_count += 1
        
        # Extract candidate text
        candidate = full_text[actual_start:actual_end]
        
        # Verify it normalizes to our search string
        if self._normalize_text(candidate) == normalized_search:
            return candidate
        
        # If exact match failed, try to find by searching for key words
        # Split normalized search into words
        search_words = normalized_search.split()
        if len(search_words) < 2:
            return None
        
        # IMPORTANT: Search in normalized text, not original text
        # This ensures we find matches regardless of case, whitespace differences, etc.
        # The normalized_full was already computed above (line 450)
        first_word = search_words[0]
        last_word = search_words[-1]
        
        # Find first word in normalized text (case-insensitive search)
        # Since _normalize_text doesn't lowercase, we need case-insensitive search
        import re
        first_match = re.search(re.escape(first_word), normalized_full, re.IGNORECASE)
        if not first_match:
            return None
        first_pos_norm = first_match.start()
        
        # Find last word after first word in normalized text (case-insensitive)
        last_match = re.search(re.escape(last_word), normalized_full[first_pos_norm:], re.IGNORECASE)
        if not last_match:
            return None
        last_pos_norm = first_pos_norm + last_match.start()
        
        # Now we need to map normalized positions back to actual text positions
        # Build a mapping by processing the text character by character
        # and tracking how normalization affects positions
        norm_to_actual_start = {}  # Maps normalized char index to actual char index
        norm_to_actual_end = {}    # Maps normalized char index to actual char end index
        
        norm_index = 0
        for actual_index, char in enumerate(full_text):
            # Normalize this single character (or sequence)
            char_normalized = self._normalize_text(char)
            
            # Map each normalized character position
            for norm_char in char_normalized:
                if norm_index not in norm_to_actual_start:
                    norm_to_actual_start[norm_index] = actual_index
                norm_to_actual_end[norm_index] = actual_index + 1
                norm_index += 1
        
        # Get actual text positions from normalized positions
        if first_pos_norm not in norm_to_actual_start:
            return None
        
        actual_start = norm_to_actual_start[first_pos_norm]
        
        # For the end, find where the last word ends in normalized text
        last_word_end_norm = last_pos_norm + len(last_word)
        if last_word_end_norm in norm_to_actual_end:
            actual_end = norm_to_actual_end[last_word_end_norm - 1]
        else:
            # Fallback: use the last available position
            actual_end = len(full_text)
        
        # Extract candidate text from actual text
        candidate = full_text[actual_start:actual_end]
        
        # Verify normalization matches (case-insensitive comparison)
        candidate_normalized = self._normalize_text(candidate)
        if candidate_normalized.lower() == normalized_search.lower():
            return candidate
        
        return None
    
    def _find_best_keyword_match(
        self,
        paragraph: Paragraph,
        keywords: list,
        normalized_search: str
    ) -> Optional[str]:
        """
        Find the best matching text in paragraph based on keywords
        Returns the actual text that best matches the search string
        """
        para_text = paragraph.text
        para_normalized = self._normalize_text(para_text)
        
        # Find positions of all keywords
        keyword_positions = []
        for keyword in keywords:
            pos = para_normalized.lower().find(keyword.lower())
            if pos != -1:
                keyword_positions.append(pos)
        
        if len(keyword_positions) < 2:
            return None
        
        # Find the span that contains all keywords
        min_pos = min(keyword_positions)
        max_pos = max(keyword_positions) + max(len(kw) for kw in keywords)
        
        # Extract candidate text (with some padding)
        # Map normalized positions back to actual text
        norm_index = 0
        actual_start = 0
        actual_end = len(para_text)
        
        for i, char in enumerate(para_text):
            char_norm = self._normalize_text(char)
            if char_norm:
                if norm_index <= min_pos < norm_index + len(char_norm):
                    actual_start = i
                if norm_index <= max_pos < norm_index + len(char_norm):
                    actual_end = i + 1
                    break
                norm_index += len(char_norm)
        
        # Extract candidate with some context
        padding = 50  # Add some context
        actual_start = max(0, actual_start - padding)
        actual_end = min(len(para_text), actual_end + padding)
        
        candidate = para_text[actual_start:actual_end]
        candidate_normalized = self._normalize_text(candidate)
        
        # Check if normalized candidate contains the search string (or most of it)
        if len(normalized_search) > 0:
            # Calculate similarity (simple word overlap)
            search_words = set(normalized_search.lower().split())
            candidate_words = set(candidate_normalized.lower().split())
            overlap = len(search_words & candidate_words)
            similarity = overlap / len(search_words) if search_words else 0
            
            # Return if similarity is high enough (at least 60% word overlap)
            if similarity >= 0.6:
                return candidate
        
        return None
    
    async def validate_document(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate that a file is a valid .docx document
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            if not file_path.endswith('.docx'):
                return False, "File is not a .docx document"
            
            # Try to open the document
            doc = Document(file_path)
            
            # Check if document has content
            has_content = False
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    has_content = True
                    break
            
            if not has_content:
                # Check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                has_content = True
                                break
                        if has_content:
                            break
                    if has_content:
                        break
            
            if not has_content:
                return False, "Document appears to be empty"
            
            return True, None
            
        except Exception as e:
            return False, f"Error reading document: {str(e)}"
    
    async def extract_text(self, file_path: str) -> List[str]:
        """
        Extract all text from a document as a list of paragraphs
        Useful for debugging and validation
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                paragraphs.append(paragraph.text)
            
            return paragraphs
            
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return []