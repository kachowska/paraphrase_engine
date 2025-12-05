"""
Block 4: Line Range Tracker
Отслеживает диапазоны строк/абзацев для фрагментов в документах
"""

import logging
from typing import List, Tuple, Optional, Dict
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType

logger = logging.getLogger(__name__)


class LineRangeTracker:
    """
    Отслеживает диапазоны строк для фрагментов текста в DOCX документах
    """
    
    def __init__(self, docx_path: str):
        """
        Инициализация трекера
        
        Args:
            docx_path: Путь к DOCX файлу
        """
        self.docx_path = docx_path
        self.doc: Optional[DocumentType] = None
        self._paragraph_map: Optional[List[Dict]] = None
        self._load_document()
        self._build_paragraph_map()
    
    def _load_document(self):
        """Загружает DOCX документ"""
        try:
            self.doc = Document(self.docx_path)
        except Exception as e:
            logger.error(f"Ошибка при загрузке документа {self.docx_path}: {e}")
            raise
    
    def _build_paragraph_map(self):
        """Строит карту параграфов документа"""
        if self.doc is None:
            self._load_document()
        
        if self.doc is None:
            return
        
        self._paragraph_map = []
        
        # Обрабатываем параграфы
        if self.doc.paragraphs:
            for para_idx, para in enumerate(self.doc.paragraphs):
                if para.text.strip():
                    self._paragraph_map.append({
                        'type': 'paragraph',
                        'index': para_idx,
                        'text': para.text,
                        'element': para
                    })
        
        # Обрабатываем таблицы
        if self.doc.tables:
            for table_idx, table in enumerate(self.doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                self._paragraph_map.append({
                                    'type': 'table',
                                    'table_index': table_idx,
                                    'row_index': row_idx,
                                    'cell_index': cell_idx,
                                    'para_index': para_idx,
                                    'text': para.text,
                                    'element': para
                                })
    
    def _find_paragraph_for_fragment(self, fragment_text: str) -> Optional[Dict]:
        """
        Находит параграф, содержащий фрагмент текста
        
        Args:
            fragment_text: Текст фрагмента для поиска
            
        Returns:
            Словарь с информацией о параграфе или None
        """
        if self._paragraph_map is None:
            self._build_paragraph_map()
        
        if self._paragraph_map is None:
            return None
        
        # Нормализуем текст фрагмента (убираем лишние пробелы)
        normalized_fragment = ' '.join(fragment_text.split())
        
        # Ищем точное совпадение или частичное
        for para_info in self._paragraph_map:
            para_text = para_info['text']
            normalized_para = ' '.join(para_text.split())
            
            # Проверяем точное совпадение
            if normalized_fragment in normalized_para:
                return para_info
            
            # Проверяем частичное совпадение (если фрагмент длиннее 50 символов)
            if len(normalized_fragment) > 50:
                # Ищем начало фрагмента в параграфе
                if normalized_para.startswith(normalized_fragment[:50]):
                    return para_info
        
        return None
    
    def _calculate_line_range(self, para_info: Dict, fragment_text: str) -> Tuple[int, int]:
        """
        Вычисляет диапазон строк для фрагмента в параграфе
        
        Args:
            para_info: Информация о параграфе
            fragment_text: Текст фрагмента
            
        Returns:
            Кортеж (start_line, end_line)
        """
        if self._paragraph_map is None:
            self._build_paragraph_map()
        
        if self._paragraph_map is None:
            return (0, 0)
        
        # Находим индекс параграфа в карте
        para_index = None
        for idx, para in enumerate(self._paragraph_map):
            if para == para_info:
                para_index = idx
                break
        
        if para_index is None:
            return (0, 0)
        
        # Вычисляем примерный диапазон строк
        # Предполагаем, что один параграф = примерно 1-3 строки
        start_line = para_index + 1
        end_line = para_index + 1
        
        # Если фрагмент длинный, увеличиваем диапазон
        if len(fragment_text) > 200:
            end_line = para_index + 2
        
        return (start_line, end_line)
    
    def find_line_range(self, fragment_text: str) -> Optional[Tuple[int, int]]:
        """
        Находит диапазон строк для фрагмента текста
        
        Args:
            fragment_text: Текст фрагмента
            
        Returns:
            Кортеж (start_line, end_line) или None если не найдено
        """
        para_info = self._find_paragraph_for_fragment(fragment_text)
        
        if para_info is None:
            logger.warning(f"Не удалось найти параграф для фрагмента: {fragment_text[:50]}...")
            return None
        
        return self._calculate_line_range(para_info, fragment_text)
    
    def get_all_fragments_with_ranges(self, fragments: List[str]) -> List[Dict]:
        """
        Получает все фрагменты с их диапазонами строк
        
        Args:
            fragments: Список текстовых фрагментов
            
        Returns:
            Список словарей с информацией о фрагментах и их диапазонах
        """
        result = []
        
        for fragment in fragments:
            line_range = self.find_line_range(fragment)
            result.append({
                'text': fragment,
                'line_range': line_range,
                'found': line_range is not None
            })
        
        return result

