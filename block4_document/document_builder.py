"""
Block 4: Document Builder
Automatically replaces original fragments with paraphrased ones in .docx documents
"""

import os
import logging
from typing import List, Tuple, Optional
from pathlib import Path
import shutil
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..block5_logging.logger import SystemLogger

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Handles document manipulation and fragment replacement"""
    
    def __init__(self):
        self.system_logger = SystemLogger()
    
    async def replace_fragments(
        self,
        source_file_path: str,
        output_file_path: str,
        original_fragments: List[str],
        paraphrased_fragments: List[str]
    ) -> bool:
        """
        Replace fragments in document
        CRITICAL: Replacements are made in reverse order to preserve text indexing
        
        Args:
            source_file_path: Path to source .docx file
            output_file_path: Path for output .docx file
            original_fragments: List of original text fragments
            paraphrased_fragments: List of paraphrased text fragments
            
        Returns:
            bool: True if successful, False otherwise
        """
        
        if len(original_fragments) != len(paraphrased_fragments):
            logger.error("Mismatch between original and paraphrased fragments count")
            return False
        
        try:
            # Create a copy of the source file
            shutil.copy2(source_file_path, output_file_path)
            logger.info(f"Created document copy: {output_file_path}")
            
            # Open the document
            doc = Document(output_file_path)
            
            # Process replacements in REVERSE order (critical requirement)
            replacements_made = 0
            skipped_fragments = []
            
            for i in range(len(original_fragments) - 1, -1, -1):
                original = original_fragments[i]
                paraphrased = paraphrased_fragments[i]
                
                logger.info(f"Processing fragment {i + 1}/{len(original_fragments)} (reverse order)")
                
                # Try to replace the fragment
                replaced = await self._replace_fragment_in_document(
                    doc,
                    original,
                    paraphrased,
                    fragment_index=i
                )
                
                if replaced:
                    replacements_made += 1
                    await self.system_logger.log_fragment_replaced(
                        fragment_index=i,
                        original_length=len(original),
                        paraphrased_length=len(paraphrased)
                    )
                else:
                    skipped_fragments.append(i)
                    await self.system_logger.log_fragment_not_found(
                        fragment_index=i,
                        fragment_text=original[:50] + "..." if len(original) > 50 else original
                    )
            
            # Save the modified document
            doc.save(output_file_path)
            
            logger.info(
                f"Document processing complete. "
                f"Replacements: {replacements_made}/{len(original_fragments)}, "
                f"Skipped: {len(skipped_fragments)}"
            )
            
            # Log completion
            await self.system_logger.log_document_processed(
                source_path=source_file_path,
                output_path=output_file_path,
                total_fragments=len(original_fragments),
                replaced_fragments=replacements_made,
                skipped_fragments=len(skipped_fragments)
            )
            
            return True
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            await self.system_logger.log_error(
                chat_id=0,
                operation="document_builder",
                error_message=str(e)
            )
            return False
    
    async def _replace_fragment_in_document(
        self,
        doc: Document,
        original_text: str,
        replacement_text: str,
        fragment_index: int
    ) -> bool:
        """
        Replace a single fragment in the document
        
        Returns:
            bool: True if replacement was made, False if fragment not found
        """
        
        # Normalize texts for comparison
        original_normalized = self._normalize_text(original_text)
        
        # Track if replacement was made
        replacement_made = False
        
        # Search through all paragraphs
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            paragraph_normalized = self._normalize_text(paragraph_text)
            
            # Check if the fragment exists in this paragraph
            if original_normalized in paragraph_normalized:
                # Find the exact position
                start_pos = paragraph_normalized.find(original_normalized)
                
                if start_pos != -1:
                    # Perform replacement while preserving formatting
                    success = self._replace_in_paragraph_with_formatting(
                        paragraph,
                        original_text,
                        replacement_text
                    )
                    
                    if success:
                        replacement_made = True
                        logger.info(f"Replaced fragment {fragment_index} in paragraph")
                        break
        
        # Also check in tables
        if not replacement_made:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text
                            paragraph_normalized = self._normalize_text(paragraph_text)
                            
                            if original_normalized in paragraph_normalized:
                                success = self._replace_in_paragraph_with_formatting(
                                    paragraph,
                                    original_text,
                                    replacement_text
                                )
                                
                                if success:
                                    replacement_made = True
                                    logger.info(f"Replaced fragment {fragment_index} in table")
                                    break
                        
                        if replacement_made:
                            break
                    
                    if replacement_made:
                        break
                
                if replacement_made:
                    break
        
        return replacement_made
    
    def _replace_in_paragraph_with_formatting(
        self,
        paragraph: Paragraph,
        original_text: str,
        replacement_text: str
    ) -> bool:
        """
        Replace text in paragraph while preserving formatting
        
        This is a complex operation that tries to maintain the original formatting
        """
        
        try:
            # Get full paragraph text
            full_text = paragraph.text
            
            # Check if original text exists
            if original_text not in full_text:
                return False
            
            # Find position of text to replace
            start_index = full_text.find(original_text)
            end_index = start_index + len(original_text)
            
            # Build new paragraph content
            new_runs = []
            current_pos = 0
            
            for run in paragraph.runs:
                run_text = run.text
                run_start = current_pos
                run_end = current_pos + len(run_text)
                
                # Case 1: Run is completely before the replacement area
                if run_end <= start_index:
                    new_runs.append((run_text, run))
                
                # Case 2: Run is completely after the replacement area
                elif run_start >= end_index:
                    new_runs.append((run_text, run))
                
                # Case 3: Run contains the start of replacement area
                elif run_start < start_index and run_end > start_index:
                    # Part before replacement
                    before_text = run_text[:start_index - run_start]
                    if before_text:
                        new_runs.append((before_text, run))
                    
                    # Add replacement text with same formatting
                    if run_end >= end_index:
                        # Replacement is completely within this run
                        new_runs.append((replacement_text, run))
                        # Part after replacement
                        after_text = run_text[end_index - run_start:]
                        if after_text:
                            new_runs.append((after_text, run))
                    else:
                        # Replacement spans multiple runs
                        new_runs.append((replacement_text, run))
                
                # Case 4: Run is completely within replacement area
                elif run_start >= start_index and run_end <= end_index:
                    # Skip this run (it's being replaced)
                    pass
                
                # Case 5: Run contains the end of replacement area
                elif run_start < end_index and run_end > end_index:
                    # Part after replacement
                    after_text = run_text[end_index - run_start:]
                    if after_text:
                        new_runs.append((after_text, run))
                
                current_pos = run_end
            
            # Clear paragraph and rebuild with new runs
            paragraph.clear()
            
            for text, original_run in new_runs:
                new_run = paragraph.add_run(text)
                # Copy formatting from original run
                self._copy_run_formatting(original_run, new_run)
            
            return True
            
        except Exception as e:
            logger.error(f"Error replacing text with formatting: {e}")
            
            # Fallback: Simple replacement without formatting preservation
            try:
                paragraph.text = full_text.replace(original_text, replacement_text)
                return True
            except:
                return False
    
    def _copy_run_formatting(self, source_run: Run, target_run: Run):
        """Copy formatting from source run to target run"""
        try:
            if source_run.bold is not None:
                target_run.bold = source_run.bold
            if source_run.italic is not None:
                target_run.italic = source_run.italic
            if source_run.underline is not None:
                target_run.underline = source_run.underline
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            if source_run.font.highlight_color:
                target_run.font.highlight_color = source_run.font.highlight_color
        except Exception as e:
            logger.debug(f"Could not copy some formatting: {e}")
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize text for comparison
        Handles different whitespace, line breaks, etc.
        """
        # Replace multiple spaces with single space
        import re
        text = re.sub(r'\s+', ' ', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        return text
    
    async def validate_document(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate that a file is a valid .docx document
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            if not file_path.endswith('.docx'):
                return False, "File is not a .docx document"
            
            # Try to open the document
            doc = Document(file_path)
            
            # Check if document has content
            has_content = False
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    has_content = True
                    break
            
            if not has_content:
                # Check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                has_content = True
                                break
                        if has_content:
                            break
                    if has_content:
                        break
            
            if not has_content:
                return False, "Document appears to be empty"
            
            return True, None
            
        except Exception as e:
            return False, f"Error reading document: {str(e)}"
    
    async def extract_text(self, file_path: str) -> List[str]:
        """
        Extract all text from a document as a list of paragraphs
        Useful for debugging and validation
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                paragraphs.append(paragraph.text)
            
            return paragraphs
            
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return []
